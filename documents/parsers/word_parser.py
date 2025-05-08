import docx
from io import BytesIO
import logging

logger = logging.getLogger(__name__)


def extract_text_from_docx(file):
    """
    Extrae texto y estructura de un archivo DOCX

    Args:
        file: Objeto archivo DOCX

    Returns:
        dict: Diccionario con texto extraído y estructura
    """
    result = {
        'text': '',
        'metadata': {},
        'paragraphs': [],
        'headings': [],
        'tables': []
    }

    try:
        # Manejar diferentes tipos de entrada
        if hasattr(file, 'read'):
            file_content = file.read()
            docx_file = BytesIO(file_content)
        else:
            docx_file = BytesIO(file)

        doc = docx.Document(docx_file)

        # Extraer metadatos
        core_properties = doc.core_properties
        result['metadata'] = {
            'author': core_properties.author,
            'title': core_properties.title,
            'created': str(core_properties.created) if core_properties.created else None,
            'modified': str(core_properties.modified) if core_properties.modified else None
        }

        # Extraer texto completo y párrafos
        full_text = ""
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                full_text += para.text + "\n"
                # Determinar si es un encabezado
                is_heading = para.style.name.startswith('Heading')
                if is_heading:
                    level = int(para.style.name.replace('Heading', '')) if para.style.name != 'Heading' else 1
                    result['headings'].append({
                        'level': level,
                        'text': para.text,
                        'position': i
                    })

                result['paragraphs'].append({
                    'text': para.text,
                    'is_heading': is_heading
                })

        # Extraer tablas
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)

            result['tables'].append({
                'position': i,
                'data': table_data
            })

        result['text'] = full_text
        return result

    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        result['error'] = str(e)
        return result